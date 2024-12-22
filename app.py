from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
from datetime import datetime
import os
import platform
import subprocess

port = int(os.environ.get("PORT", 8501))
# Path to the text file for storing base number and counter
SERIAL_FILE = "serial_data.txt"

def get_serial_number():
    # Read base number and counter from file
    with open(SERIAL_FILE, "r") as f:
        base_number, counter = map(int, f.read().strip().split(","))

    # Calculate current serial number
    serial_number = base_number + counter

    # Increment the counter and update the file
    with open(SERIAL_FILE, "w") as f:
        f.write(f"{base_number},{counter + 1}")

    return serial_number

def generate_reference_number(company_name="BKR"):
    """
    Generate the full reference number in the format: BKRMM-YYYY-CR<serial>.
    """
    current_month = datetime.now().strftime("%m")
    current_year = datetime.now().strftime("%Y")
    serial_number = get_serial_number()
    return f"{company_name}{current_month}-{current_year}-CR{serial_number}"

# Code to replace placeholder for VAT
def replace_placeholders_vat(doc, placeholders):
    """Replace placeholders in a Word document, maintaining original formatting."""
    
    def replace_in_paragraph(paragraph, key, value):
        """Replace placeholders in a paragraph, preserving formatting."""
        for run in paragraph.runs:
            if key in run.text:
                # Replace placeholder text
                run.text = run.text.replace(key, value)
                # Retain original font style and size
                run.font.name = paragraph.style.font.name
                run.font.size = paragraph.style.font.size

    def replace_in_cell(cell, placeholders):
        """Replace placeholders inside a table cell."""
        for para in cell.paragraphs:
            for key, value in placeholders.items():
                replace_in_paragraph(para, key, value)

    # Replace placeholders in all paragraphs
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            replace_in_paragraph(para, key, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, placeholders)

    return doc

def replace_placeholders(doc, placeholders):
    """Replace placeholders in a Word document, including paragraphs and tables."""
    
    def replace_in_paragraph(paragraph, key, value):
        """Replace placeholders in a single paragraph, handling split runs."""
        full_text = "".join(run.text for run in paragraph.runs)
        if key in full_text:
            full_text = full_text.replace(key, value)
            for run in paragraph.runs:
                run.text = ""  # Clear all runs
            paragraph.runs[0].text = full_text  # Add the replaced text back

    def replace_in_cell(cell, placeholders):
        """Replace placeholders inside a table cell."""
        for para in cell.paragraphs:
            for key, value in placeholders.items():
                replace_in_paragraph(para, key, value)

    # Replace placeholders in all paragraphs
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            replace_in_paragraph(para, key, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, placeholders)

    return doc
  


def convert_to_pdf(doc_path, pdf_path):
    doc_path = os.path.abspath(doc_path)
    pdf_path = os.path.abspath(pdf_path)

    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Word document not found at {doc_path}")

    if platform.system() == "Windows":
        try:
            import comtypes.client
            import pythoncom
            pythoncom.CoInitialize()
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(doc_path)
            doc.SaveAs(pdf_path, FileFormat=17)
            doc.Close()
            word.Quit()
        except Exception as e:
            raise Exception(f"Error using COM on Windows: {e}")
    else:
        try:
            subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), doc_path],
                check=True
            )
        except subprocess.CalledProcessError as e:
            raise Exception(f"Error using LibreOffice: {e}")

def options_changed():
    if "current_input" not in st.session_state:
        return False
    return st.session_state["current_input"] != current_input

def generate_unique_reference():
    """
    Generate a unique reference number based on the current date and time in the format:
    DDMMYYYYHHMMSS
    """
    now = datetime.now()
    return now.strftime("%d%m%Y%H%M%S")

st.title("Generator")

# VAT Registration Fields
template_option = st.selectbox("Select Template", ["VAT Registration", "Service Agreement","Invoice"])
current_input = {}

if template_option == "VAT Registration":
    agreement_date = st.date_input("Date of Agreement", datetime.today())
    atten = st.text_input("Attention")
    email = st.text_input("Email")
    client_name = st.text_input("Client Name")
    commercial_registration_number = st.text_input("Commercial Registration Number")
    service_provider_name = st.text_input("Service Provider Name")
    service_provider_cr = st.text_input("Service Provider CR Number")
    company_name = st.text_input("Company Name")
    vat_registration_fee = st.text_input("VAT Registration Fee")
    consultancy_fee = st.text_input("Consultancy Fee")
    authorized_person_name = st.text_input("Authorized Person Name")

    # Prepare inputs for comparison
    current_input = {
        "template": template_option,
        "agreement_date": agreement_date,
        "atten": atten,
        "email": email,
        "client_name": client_name,
        "commercial_registration_number": commercial_registration_number,
        "service_provider_name": service_provider_name,
        "service_provider_cr": service_provider_cr,
        "company_name": company_name,
        "vat_registration_fee": vat_registration_fee,
        "consultancy_fee": consultancy_fee,
        "authorized_person_name": authorized_person_name,
    }

    if st.button("Generate VAT Document"):
        if not all(current_input.values()) :
            st.error("Please fill all fields and upload the signature image!")
        else:
            try:
                reference_number = generate_reference_number()
                placeholders = {
                    "<<Date>>": agreement_date.strftime("%d-%m-%Y"),
                    "<<Atten>>": atten,
                    "<<Email>>": email,
                    "<<Client Name>>": client_name,
                    "<<Commercial Registration Number>>": commercial_registration_number,
                    "<<Service Provider CR Number>>": service_provider_cr,
                    "<<Company Name>>": company_name,
                    "<<VAT Registration Fee>>": vat_registration_fee,
                    "<<Consultancy Fee>>": consultancy_fee,
                    "<<Authorized Person Name>>": authorized_person_name,
                    "<<Reference Number>>": reference_number,
                    "<<Service Provider Name>>":service_provider_name
                }

                # Generate document
                template_path = "SAMPLE VAT registration and VAT filling -SME package.docx"
                doc = Document(template_path)
                doc = replace_placeholders_vat(doc, placeholders)

                word_output = f"VAT {client_name}.docx"
                pdf_output = word_output.replace(".docx", ".pdf")
                doc.save(word_output)
                convert_to_pdf(word_output, pdf_output)

                st.success("Document generated successfully!")
                st.session_state["current_input"] = current_input
                st.session_state["word_output"] = word_output
                st.session_state["pdf_output"] = pdf_output

            except Exception as e:
                st.error(f"Error: {e}")

    # Display download buttons only if options haven't changed
    if not options_changed() and "word_output" in st.session_state and "pdf_output" in st.session_state:
        with open(st.session_state["word_output"], "rb") as word_file:
            st.download_button("Download VAT Document (Word)", word_file, file_name=st.session_state["word_output"])
        with open(st.session_state["pdf_output"], "rb") as pdf_file:
            st.download_button("Download VAT Document (PDF)", pdf_file, file_name=st.session_state["pdf_output"])

    elif options_changed():
        st.session_state.pop("word_output", None)
        st.session_state.pop("pdf_output", None)

elif template_option == "Service Agreement":
    agreement_date = st.date_input("Date of Agreement", datetime.today())
    client_name = st.text_input("Client Name")
    bahraini_ownership = st.number_input("Bahraini Ownership (%)", min_value=0, max_value=100, step=1)
    gcc_ownership = st.number_input("GCC Nationals Ownership (%)", min_value=0, max_value=100, step=1)
    american_ownership = st.number_input("American Nationals Ownership (%)", min_value=0, max_value=100, step=1)
    foreign_ownership = st.number_input("Foreign Ownership (%)", min_value=0, max_value=100, step=1)
    business_activity_1_isic = st.text_input("Business Activity ISIC4 Code (1st)")
    business_activity_1_name = st.text_input("Business Activity Name (1st)")
    business_activity_1_desc = st.text_area("Business Activity Description (1st)")
    business_activity_2_isic = st.text_input("Business Activity ISIC4 Code (2nd)")
    business_activity_2_name = st.text_input("Business Activity Name (2nd)")
    business_activity_2_desc = st.text_area("Business Activity Description (2nd)")
    costs = {
        "Company Formation Cost": st.number_input("Company Formation Cost", min_value=0.0, step=0.01),
        "Desk-Space Office Rental Cost": st.number_input("Desk-Space Office Rental Cost", min_value=0.0, step=0.01),
        "Businessman Visa Cost": st.number_input("Businessman Visa Cost", min_value=0.0, step=0.01),
        "Miscellaneous/Admin Charges": st.number_input("Miscellaneous/Admin Charges", min_value=0.0, step=0.01),
        "Power of Attorney Cost": st.number_input("Power of Attorney Cost", min_value=0.0, step=0.01),
        "Estimation Charges (Per Head)": st.number_input("Estimation Charges (Per Head)", min_value=0.0, step=0.01),
        "Labour Authority Registration Cost": st.number_input("Labour Authority Registration Cost", min_value=0.0, step=0.01),
        "Social Insurance Registration Cost": st.number_input("Social Insurance Registration Cost", min_value=0.0, step=0.01),
        "Free Advice/Guidance Cost": st.number_input("Free Advice/Guidance Cost", min_value=0.0, step=0.01),
    }
    total_cost = sum(costs.values())

    signatory_name = st.text_input("Signatory Name")
    passport_number = st.text_input("Passport Number")

    current_input = {
        "template": template_option,
        "agreement_date": agreement_date,
        "client_name": client_name,
        "bahraini_ownership": bahraini_ownership,
        "gcc_ownership": gcc_ownership,
        "american_ownership": american_ownership,
        "foreign_ownership": foreign_ownership,
        "business_activity_1_isic": business_activity_1_isic,
        "business_activity_1_name": business_activity_1_name,
        "business_activity_1_desc": business_activity_1_desc,
        "business_activity_2_isic": business_activity_2_isic,
        "business_activity_2_name": business_activity_2_name,
        "business_activity_2_desc": business_activity_2_desc,
        "company_formation_cost": costs["Company Formation Cost"],
        "desk_rental_cost": costs["Desk-Space Office Rental Cost"],
        "businessman_visa_cost": costs["Businessman Visa Cost"],
        "misc_admin_charges": costs["Miscellaneous/Admin Charges"],
        "power_of_attorney_cost": costs["Power of Attorney Cost"],
        "estimation_charges": costs["Estimation Charges (Per Head)"],
        "labour_registration_cost": costs["Labour Authority Registration Cost"],
        "social_insurance_cost": costs["Social Insurance Registration Cost"],
        "free_advice_cost": costs["Free Advice/Guidance Cost"],
        "total_cost": total_cost,
        "signatory_name": signatory_name,
        "passport_number": passport_number,
       
    }


    if st.button("Generate Service Agreement Document"):
        if not all([
            client_name.strip(),
            signatory_name.strip(),
            passport_number.strip(),
            business_activity_1_isic,
            business_activity_1_name ,
            business_activity_1_desc, 
            business_activity_2_isic ,
            business_activity_2_name,
            business_activity_2_desc 
        ]):
            st.error("Please fill all required fields!")
        else:
            try:
                reference_number = generate_reference_number()
                placeholders = {
    "<< Date >>": agreement_date.strftime("%d-%m-%Y"),
    "<< Reference Number >>": reference_number,
    "<< Client Name >>": client_name,
    "<< Bahraini Ownership >>": f"{bahraini_ownership}%",
    "<< GCC Nationals Ownership >>": f"{gcc_ownership}%",
    "<< American Nationals Ownership >>": f"{american_ownership}%",
    "<< Foreign Ownership >>": f"{foreign_ownership}%",
    "<<Text1>>": business_activity_1_isic,
    "<<Text2>>": business_activity_1_name,
    "<<Text3>>": business_activity_1_desc,
    "<<Text4>>": business_activity_2_isic,
    "<<Text5>>": business_activity_2_name,
    "<<Text6>>": business_activity_2_desc,
    "<< Company Formation Cost >>": f"{costs['Company Formation Cost']:.2f}",
    "<< Desk-Space Office Rental Cost >>": f"{costs['Desk-Space Office Rental Cost']:.2f}",
    "<< Businessman Visa Cost >>": f"{costs['Businessman Visa Cost']:.2f}",
    "<< Miscellaneous/Admin Charges >>": f"{costs['Miscellaneous/Admin Charges']:.2f}",
    "<< Power of Attorney Cost >>": f"{costs['Power of Attorney Cost']:.2f}",
    "<< Estimation Charges (Per Head) >>": f"{costs['Estimation Charges (Per Head)']:.2f}",
    "<< Labour Authority Registration Cost >>": f"{costs['Labour Authority Registration Cost']:.2f}",
    "<< Social Insurance Registration Cost >>": f"{costs['Social Insurance Registration Cost']:.2f}",
    "<< Free Advice/Guidance Cost >>": f"{costs['Free Advice/Guidance Cost']:.2f}",
    "<< Total Cost >>": f"{total_cost:.2f}",
    "<< Signatory Name >>": signatory_name,
    "<< Passport Number >>": passport_number,
}
                
                doc = replace_placeholders(Document("SAMPLE Service Agreement -Company formation -Bahrain - Filled.docx"), placeholders)


                word_output = f"Service Agreement {client_name}.docx"
                pdf_output = word_output.replace(".docx", ".pdf")
                doc.save(word_output)
                convert_to_pdf(word_output, pdf_output)

                st.success("Document generated successfully!")
                st.session_state["current_input"] = current_input
                st.session_state["word_output"] = word_output
                st.session_state["pdf_output"] = pdf_output

            except Exception as e:
                st.error(f"Error: {e}")

    if not options_changed() and "word_output" in st.session_state and "pdf_output" in st.session_state:
        with open(st.session_state["word_output"], "rb") as word_file:
            st.download_button("Download Service Agreement (Word)", word_file, file_name=st.session_state["word_output"])
        with open(st.session_state["pdf_output"], "rb") as pdf_file:
            st.download_button("Download Service Agreement (PDF)", pdf_file, file_name=st.session_state["pdf_output"])
    
    elif options_changed():
        st.session_state.pop("word_output", None)
        st.session_state.pop("pdf_output", None)
        
        
elif template_option == "Invoice":
    
    service_data = {
        "LMRA Affairs": [
            "Visa Application", "Visa Termination", "Visa Renewal", "Visa Ceiling Application", "Changing Occupation", "Mobility Issues", "Offences Removal Application", "Runaway Application", "Domestic Permit Application", "LMRA Registration of Establishments", "Work Load Application", "Biometrics Appointment"
        ],
        "NPRA (Immigration) Affairs": [
            "Visa Cancellation and Extension", "Dependent Visa Processing", "Domestic Visa Processing", "Visit Visa Extension", "Business Visit Visa Processing", "Dependent Visit Visa Processing", "Visa Cancellation Update", "Passport Update", "RP Stamping", "eVisa Processing", "Business Investor Visa Processing"
        ],
        "SIO Affairs": [
            "Employee's Registration", "Employee's Termination", "Payment Processing", "Establishment Registration", "Addition Bahraini Employee"
        ],
        "CIO Affairs": [
            "CPR Issuance", "CPR Renewal", "CPR Update", "Dependent CPR", "Lost CPR", "Address Update"
        ],
        "CID Affairs": [
            "Report Issuance for Lost Passport", "Good Conduct Certificate Issuance", "Other Kind of Reports", "CPR Offense Inquiry and Removal"
        ],
        "eGovernment": ["Driving School Appointments","EWA Bills","Traffic Contraventions Details","Vehicle Details","Online Appointments"],
    }

    service_type = st.selectbox("Select Service Type", list(service_data.keys()))
    service = st.selectbox("Select Service", service_data[service_type])
    
    # Input Fields for Invoice
    invoice_date = st.date_input("Date", datetime.today())
    client_name = st.text_input("Client Name")
    reference_number=st.text_input("Service Agreement Reference Number")
    attention = st.text_input("Attention (Atten)")
    cost = st.text_input("Cost (in BHD)")
    total_in_words = st.text_input("Total Amount (in words)")
    total_amount = st.text_input("Total Amount (in BHD)")


    current_input = {
        "template": template_option,
        "invoice_date": invoice_date,
        "client_name": client_name,
        "attention": attention,
        "cost": cost,
        "total_in_words": total_in_words,
        "total_amount": total_amount,
        "reference_number":reference_number,
        "service": service,
        "service_type": service_type,
    }

    if st.button("Generate Invoice"):
        if not all(current_input.values()):
            st.error("Please fill all fields!")
        else:
            try:
                invoice_number=generate_unique_reference();
                placeholders = {
                    "<<Date>>": invoice_date.strftime("%d-%m-%Y"),
                    "<<Invoice Number>>": invoice_number,
                    "<<Client Name>>": client_name,
                    "<<Atten>>": attention,
                    "<<Service Agreement Ref Number>>": reference_number,
                    "<<Cost>>": cost,
                    "<<Service>>": service,
                    "<<Total In Words>>": total_in_words,
                    "<<Total Amount>>": total_amount,
                }

                template_path = "SAMPLE -Invoice BKR2024CF158 - first payment.docx"
                if not os.path.exists(template_path):
                    st.error(f"Template file not found: {template_path}")
                    raise FileNotFoundError(f"Template file not found: {template_path}")

                doc = Document(template_path)
                doc = replace_placeholders(doc, placeholders)

                word_output = f"Invoice {client_name}.docx"
                pdf_output = word_output.replace(".docx", ".pdf")

                doc.save(word_output)
                convert_to_pdf(word_output, pdf_output)

                st.success("Invoice generated successfully!")
                st.session_state["current_input"] = current_input
                st.session_state["word_output"] = word_output
                st.session_state["pdf_output"] = pdf_output

            except Exception as e:
                st.error(f"Error: {e}")

    # Display download buttons if options haven't changed
    if not options_changed() and "word_output" in st.session_state and "pdf_output" in st.session_state:
        with open(st.session_state["word_output"], "rb") as word_file:
            st.download_button("Download Invoice (Word)", word_file, file_name=os.path.basename(st.session_state["word_output"]))
        with open(st.session_state["pdf_output"], "rb") as pdf_file:
            st.download_button("Download Invoice (PDF)", pdf_file, file_name=os.path.basename(st.session_state["pdf_output"]))

    elif options_changed():
        st.session_state.pop("word_output", None)
        st.session_state.pop("pdf_output", None)